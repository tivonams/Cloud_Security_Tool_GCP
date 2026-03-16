from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_structured_answer(doc: Document, answer_text: str):
    """
    Parses LLM response and converts markdown-style headings
    into real DOCX headings.
    """
    lines = answer_text.split("\n")

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Convert #### headings to DOCX Heading level 3 (Light Gray)
        if line.startswith("#### "):
            heading_text = line.replace("#### ", "").strip()
            h = doc.add_heading(heading_text, level=3)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        # Convert ### headings to DOCX Heading level 2
        elif line.startswith("### "):
            heading_text = line.replace("### ", "").strip()
            doc.add_heading(heading_text, level=2)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(8)

def build_docx(llm_response: dict) -> Document:
    doc = Document()

    # -------------------------
    # TITLE
    # -------------------------

    title = doc.add_heading(
        "Tivona Cloud Security Analysis Report",
        level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacing

    # -------------------------
    # ANALYSIS SUMMARY
    # -------------------------
    summary_heading = doc.add_heading("Analysis Summary", level=1)
    summary_heading.bold = True

    answer_text = llm_response.get("answer", "N/A")
    add_structured_answer(doc, answer_text)

    # -------------------------
    # SOURCES
    # -------------------------
    doc.add_paragraph("")  # spacing
    sources_heading = doc.add_heading("Sources", level=1)
    sources_heading.bold = True

    sources = llm_response.get("sources", [])

    if not sources:
        doc.add_paragraph("No sources available.")
    else:
        seen = set()
        for s in sources:
            key = (s.get("title"), s.get("source"))
            if key in seen:
                continue
            seen.add(key)

            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(
                f"{s.get('title', 'Unknown')} "
                f"({s.get('provider', 'N/A')})\n"
                f"{s.get('source', '')}"
            )
            run.font.size = Pt(10)

    # -------------------------
    # FOOTER NOTE (OPTIONAL)
    # -------------------------
    doc.add_paragraph("")
    note = doc.add_paragraph(
        "This report is generated using retrieved cloud security documentation. "
        "Recommendations are advisory and must be reviewed before implementation."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)

    return doc


# import os
# from datetime import datetime
# from django.conf import settings


# def build_Doc(llm_response: dict) -> str:
#     lines = []

#     lines.append("#Tivona Cloud Security Analysis Report\n")

#     lines.append("")

#     lines.append(llm_response.get("answer", "N/A"))
#     lines.append("")

#     lines.append("## Sources")
#     sources = llm_response.get("sources", [])

#     if not sources:
#         lines.append("No sources available.")
#     else:
#         seen = set()
#         for s in sources:
#             key = (s.get("title"), s.get("source"))
#             if key in seen:
#                 continue
#             seen.add(key)

#             lines.append(
#                 f"- **{s.get('title', 'Unknown')}** ({s.get('provider', 'N/A')})  \n"
#                 f"  {s.get('source')}"
#             )

#     return "\n".join(lines)


# def save_Doc(llm_response: dict) -> dict:
#     content = build_Doc(llm_response)

#     reports_dir = os.path.join(settings.MEDIA_ROOT, "reports")
#     os.makedirs(reports_dir, exist_ok=True)

#     filename = f"rag_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
#     file_path = os.path.join(reports_dir, filename)

#     with open(file_path, "w", encoding="utf-8") as f:
#         f.write(content)

#     return {
#         "filename": filename,
#         "path": file_path,
#         "url": f"{settings.MEDIA_URL}reports/{filename}"
#     }